from pathlib import Path
from typing import Dict, Iterable, List, Tuple

from docx import Document
from docx.text.paragraph import Paragraph

from .replacement import ReplacementResolver


def _set_text_preserve_formatting(paragraph: Paragraph, text: str) -> None:
    """Replace visible paragraph text while retaining first-run formatting."""

    if not paragraph.runs:
        paragraph.add_run(text)
        return

    template = paragraph.runs[0]
    bold = template.bold
    italic = template.italic
    underline = template.underline
    font_name = template.font.name
    font_size = template.font.size
    color = (
        template.font.color.rgb
        if template.font.color and template.font.color.type
        else None
    )

    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if color:
        run.font.color.rgb = color


def resolve_non_overlapping(detections: Iterable[Dict]) -> List[Dict]:
    """Keep the strongest non-overlapping detection for each text span."""

    candidates = [
        detection
        for detection in detections
        if detection.get("start", 0) < detection.get("end", 0)
    ]

    ranked = sorted(
        candidates,
        key=lambda detection: (
            -float(detection.get("confidence", 0.0)),
            -(detection["end"] - detection["start"]),
            detection["start"],
        ),
    )

    accepted = []

    for candidate in ranked:
        overlaps = any(
            candidate["start"] < existing["end"]
            and existing["start"] < candidate["end"]
            for existing in accepted
        )

        if not overlaps:
            accepted.append(candidate)

    return sorted(
        accepted,
        key=lambda detection: (
            detection["start"],
            detection["end"],
        ),
    )


def apply_redactions(
    text: str,
    detections: Iterable[Dict],
    resolver: ReplacementResolver,
) -> Tuple[str, List[Dict]]:
    """Apply replacements from right to left so source offsets remain valid."""

    accepted = resolve_non_overlapping(detections)
    redacted = text
    audit_records = []

    for detection in reversed(accepted):
        start = detection["start"]
        end = detection["end"]
        replacement = resolver.replacement_for(
            detection["entity_type"],
            detection["text"],
        )

        redacted = redacted[:start] + replacement + redacted[end:]

        audit_record = dict(detection)
        audit_record["replacement"] = replacement
        audit_records.append(audit_record)

    return redacted, list(reversed(audit_records))


def redact_docx(
    source_path: Path,
    output_path: Path,
    redacted_by_element: Dict[str, str],
) -> None:
    """Write a redacted DOCX while leaving the source document unchanged."""

    document = Document(source_path)

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        element_id = f"P_{paragraph_index:05d}"
        if element_id in redacted_by_element:
            _set_text_preserve_formatting(
                paragraph,
                redacted_by_element[element_id],
            )

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                element_id = (
                    f"T_{table_index:04d}_"
                    f"R_{row_index:04d}_"
                    f"C_{cell_index:04d}"
                )

                if element_id in redacted_by_element:
                    cell_paragraphs = cell.paragraphs
                    _set_text_preserve_formatting(
                        cell_paragraphs[0],
                        redacted_by_element[element_id],
                    )
                    for extra_paragraph in cell_paragraphs[1:]:
                        extra_paragraph.clear()

    for section_index, section in enumerate(document.sections):
        for location, container in (
            ("header", section.header),
            ("footer", section.footer),
        ):
            for paragraph_index, paragraph in enumerate(container.paragraphs):
                element_id = (
                    f"{location[0].upper()}_"
                    f"{section_index:04d}_"
                    f"P_{paragraph_index:04d}"
                )

                if element_id in redacted_by_element:
                    _set_text_preserve_formatting(
                        paragraph,
                        redacted_by_element[element_id],
                    )

    # Core properties are not visible in the document body, but they can
    # still carry personal metadata. Clear or replace only fields captured
    # by the ingestion layer.
    core_properties = document.core_properties
    metadata_replacements = {
        "title": "PIIShield Redacted Document",
        "subject": "Redacted enterprise data document",
        "author": "PIIShield",
        "keywords": "redacted, pii",
        "comments": "Generated by PIIShield",
        "last_modified_by": "PIIShield",
    }

    for field, replacement in metadata_replacements.items():
        setattr(core_properties, field, replacement)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

