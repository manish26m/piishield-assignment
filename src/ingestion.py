
from pathlib import Path
from docx import Document
import json


def extract_document(docx_path: str):
    """
    Extract paragraphs and tables from a DOCX document
    into a structured Bronze-layer representation.
    """

    doc = Document(docx_path)

    elements = []

    # --------------------------------------------------
    # 1. Extract paragraphs
    # --------------------------------------------------
    for paragraph_index, paragraph in enumerate(doc.paragraphs):

        text = paragraph.text.strip()

        if not text:
            continue

        elements.append({
            "element_id": f"P_{paragraph_index:05d}",
            "element_type": "paragraph",
            "index": paragraph_index,
            "text": text
        })

    # --------------------------------------------------
    # 2. Extract tables
    # --------------------------------------------------
    for table_index, table in enumerate(doc.tables):

        for row_index, row in enumerate(table.rows):

            for cell_index, cell in enumerate(row.cells):

                text = cell.text.strip()

                if not text:
                    continue

                elements.append({
                    "element_id": (
                        f"T_{table_index:04d}_"
                        f"R_{row_index:04d}_"
                        f"C_{cell_index:04d}"
                    ),
                    "element_type": "table_cell",
                    "table_index": table_index,
                    "row_index": row_index,
                    "cell_index": cell_index,
                    "text": text
                })

    # --------------------------------------------------
    # 3. Extract headers and footers
    # --------------------------------------------------
    for section_index, section in enumerate(doc.sections):

        for location, container in (
            ("header", section.header),
            ("footer", section.footer),
        ):
            for paragraph_index, paragraph in enumerate(container.paragraphs):

                text = paragraph.text.strip()

                if not text:
                    continue

                elements.append({
                    "element_id": (
                        f"{location[0].upper()}_"
                        f"{section_index:04d}_"
                        f"P_{paragraph_index:04d}"
                    ),
                    "element_type": location,
                    "section_index": section_index,
                    "index": paragraph_index,
                    "text": text,
                })

    # --------------------------------------------------
    # 4. Extract core document metadata
    # --------------------------------------------------
    core_properties = doc.core_properties
    metadata_fields = (
        "title",
        "subject",
        "author",
        "keywords",
        "comments",
        "last_modified_by",
    )

    for field in metadata_fields:
        value = getattr(core_properties, field, None)
        text = str(value).strip() if value is not None else ""

        if not text:
            continue

        elements.append({
            "element_id": f"M_{field.upper()}",
            "element_type": "metadata",
            "metadata_field": field,
            "text": text,
        })

    return elements


def save_bronze(elements, output_path: str):

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with open(output_path, "w", encoding="utf-8") as f:

        for element in elements:
            f.write(json.dumps(element, ensure_ascii=False) + "\n")


if __name__ == "__main__":

    project_root = Path(__file__).resolve().parent.parent

    input_file = (
        project_root
        / "docs"
        / "Red Herring Prospectus.docx"
    )

    output_file = (
        project_root
        / "data"
        / "bronze"
        / "document_elements.jsonl"
    )

    print("Starting DOCX ingestion...")
    print(f"Input: {input_file}")

    elements = extract_document(str(input_file))

    save_bronze(elements, str(output_file))

    print()
    print("Ingestion complete.")
    print(f"Elements extracted: {len(elements)}")
    print(f"Bronze output: {output_file}")