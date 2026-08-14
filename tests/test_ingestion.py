
from docx import Document

from src.ingestion import extract_document


def test_ingestion_extracts_headers_footers_and_metadata(tmp_path):
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Body paragraph")
    document.sections[0].header.paragraphs[0].text = "Header text"
    document.sections[0].footer.paragraphs[0].text = "Footer text"
    document.core_properties.author = "Document Author"
    document.save(source)

    records = extract_document(str(source))
    by_type = {record["element_type"] for record in records}

    assert "paragraph" in by_type
    assert "header" in by_type
    assert "footer" in by_type
    assert "metadata" in by_type
    assert any(record["text"] == "Header text" for record in records)
    assert any(record["text"] == "Footer text" for record in records)
    assert any(
        record.get("metadata_field") == "author"
        for record in records
    )